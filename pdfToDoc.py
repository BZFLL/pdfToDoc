import os
import gc
import sys
import time
import json
import logging
import threading
import subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import re
import requests
from PIL import Image, ImageEnhance
import numpy as np
from pdf2image import convert_from_path
from paddleocr import PaddleOCR
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import shutil
import queue

# 配置日志
logging.basicConfig(
    filename='app.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# 从配置文件中读取API设置
with open('config.json') as f:
    config = json.load(f)
API_ENDPOINT = config['api_config']['endpoint']
AI_API_KEY = config['api_config']['key']
REQUEST_TIMEOUT = 30    # API请求超时时间

class APIRateLimiter:
    """API速率限制器"""
    def __init__(self, calls_per_minute=30):
        self.interval = 60 / calls_per_minute
        self.last_call = 0

    def wait(self):
        elapsed = time.time() - self.last_call
        if elapsed < self.interval:
            time.sleep(self.interval - elapsed)
        self.last_call = time.time()

def ocr_text(image_path):
    """使用PaddleOCR进行文字识别（不进行方向识别）"""
    try:
        if not hasattr(ocr_text, "ocr_engine"):
            ocr_text.ocr_engine = PaddleOCR(
                lang='en',
                use_gpu=False,
                total_process=4
            )
        result = ocr_text.ocr_engine.ocr(image_path, cls=False)
        if not result:
            return ""
        recognized_text = "\n".join([line[1][0] for group in result for line in group])
        return recognized_text
    except Exception as e:
        messagebox.showerror("错误", f"OCR识别失败: {str(e)}")
        logging.error(f"OCR识别失败: {str(e)}")
        return None

def split_text_into_chunks(text, max_length=1000):
    sentences = re.split(r'(?<=[。.!?])', text)
    chunks = []
    current_chunk = ""
    for sentence in sentences:
        if len(current_chunk) + len(sentence) > max_length:
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                chunks.append(sentence.strip())
                current_chunk = ""
        else:
            current_chunk += sentence
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

def translate_text(text, target_language):
    try:
        max_retries = 3
        timeout = 60
        url = API_ENDPOINT
        headers = {
            "Authorization": f"Bearer {AI_API_KEY}",
            "Content-Type": "application/json",
            "accept": "application/json"
        }
        payload = {
            "model": "deepseek-ai/DeepSeek-V3",
            "messages": [
                {
                    "role": "system",
                    "content": ("你是一个中英文翻译专家，将用户输入的中文翻译成英文，"
                                "或将用户输入的英文翻译成中文。对于非中文内容，它将提供中文翻译结果。"
                                "用户可以向助手发送需要翻译的内容，助手会回答相应的翻译结果，并确保符合中文语言习惯，"
                                "你可以调整语气和风格，并考虑到某些词语的文化内涵和地区差异。"
                                "同时作为翻译家，需将原文翻译成具有信达雅标准的译文。"
                                "\"信\" 即忠实于原文的内容与意图；"
                                "\"达\" 意味着译文应通顺易懂，表达清晰；"
                                "\"雅\" 则追求译文的文化审美和语言的优美。"
                                "目标是创作出既忠于原作精神，又符合目标语言文化和读者审美的翻译。")
                },
                {
                    "role": "user",
                    "content": f"请专业准确地翻译成{target_language}，保留所有数字和格式：\n{text}"
                }
            ],
            "stream": True,
            "max_tokens": 512,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "frequency_penalty": 0.5,
            "n": 1
        }
        for attempt in range(max_retries):
            try:
                response = requests.post(url, headers=headers, json=payload, stream=True, timeout=timeout)
                if response.status_code == 200:
                    aggregated_text = ""
                    for line in response.iter_lines():
                        if line:
                            decoded_line = line.decode('utf-8').strip()
                            if decoded_line == "[DONE]":
                                continue
                            if decoded_line.startswith("data:"):
                                decoded_line = decoded_line[5:].strip()
                            if not decoded_line or decoded_line[0] != '{':
                                logging.warning(f"跳过不符合格式的chunk: {decoded_line}")
                                continue
                            try:
                                data = json.loads(decoded_line)
                                delta = data.get("choices", [{}])[0].get("delta", {})
                                content = delta.get("content")
                                if content:
                                    aggregated_text += content
                            except Exception as e:
                                logging.error(f"解析翻译chunk出错: {e}")
                    return aggregated_text
                else:
                    messagebox.showerror("错误", f"翻译请求失败，状态码：{response.status_code}")
                    logging.error(f"翻译请求失败，状态码：{response.status_code}")
                    return None
            except requests.exceptions.Timeout:
                if attempt < max_retries - 1:
                    logging.warning(f"翻译超时，正在重试 ({attempt+1}/{max_retries})...")
                    continue
                else:
                    messagebox.showerror("错误", "翻译服务响应超时")
                    logging.error("翻译服务响应超时")
                    return None
        return None
    except Exception as e:
        messagebox.showerror("错误", f"翻译过程中出错: {str(e)}")
        logging.error(f"翻译过程中出错: {str(e)}")
        return None

def translate_text_in_chunks(text, target_language, delay=0.5):
    chunks = split_text_into_chunks(text, max_length=1000)
    translated_chunks = []
    for i, chunk in enumerate(chunks, 1):
        time.sleep(delay)
        translated = translate_text(chunk, target_language)
        if translated is None:
            translated_chunks.append("（翻译失败）")
        else:
            translated_chunks.append(translated)
        logging.info(f"第 {i} 块翻译完成，共 {len(chunks)} 块")
    return "\n".join(translated_chunks)

def pdf_to_images(pdf_path, output_folder):
    try:
        images = convert_from_path(
            pdf_path, 
            dpi=config['ocr_settings']['dpi'],
            poppler_path=config['poppler_config']['path']
        )
    except Exception as e:
        messagebox.showerror("错误", f"PDF转换失败: {str(e)}")
        logging.error(f"PDF转换失败: {str(e)}")
        return []
    image_files = []
    for i, image in enumerate(images):
        image_path = os.path.join(output_folder, f"temp_page_{i + 1}.png")
        image.save(image_path, 'PNG')
        image.close()
        image_files.append(image_path)
    return image_files

def enhance_image(image_path, output_path, enhance_level=1.5):
    image = Image.open(image_path)
    enhancer = ImageEnhance.Sharpness(image)
    enhanced_image = enhancer.enhance(enhance_level)
    enhanced_image.save(output_path)

class ImageProcessorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF智能转换器")
        self.target_language = 'ch'  # 默认翻译成中文
        
        # macOS特定的窗口设置
        self.root.configure(bg='#E8E8E8')
        self.root.geometry("800x600")
        
        # 设置窗口样式
        style = ttk.Style()
        style.configure('TButton', padding=10)
        style.configure('TLabel', padding=10)
        
        # 创建主框架
        main_frame = ttk.Frame(root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 标题
        title_label = ttk.Label(main_frame, text="PDF智能转换系统", font=('Helvetica', 20))
        title_label.grid(row=0, column=0, pady=20)
        
        # 语言选择
        language_label = ttk.Label(main_frame, text="选择翻译语言", font=('Helvetica', 12))
        language_label.grid(row=1, column=0, pady=10)
        
        self.language_var = tk.StringVar(value="ch")
        language_combobox = ttk.Combobox(main_frame, textvariable=self.language_var, values=['en', 'ch', 'es', 'fr'])
        language_combobox.grid(row=2, column=0, pady=10)
        
        # 选择文件按钮
        self.folder_button = ttk.Button(main_frame, text="选择PDF文件", command=self.select_pdf_file)
        self.folder_button.grid(row=3, column=0, pady=10)
        
        # 显示选择的文件路径
        self.folder_label = ttk.Label(main_frame, text="未选择文件", font=('Helvetica', 12))
        self.folder_label.grid(row=4, column=0, pady=10)
        
        # 图片大小选择
        size_frame = ttk.LabelFrame(main_frame, text="图片大小", padding="10")
        size_frame.grid(row=5, column=0, pady=20, sticky=(tk.W, tk.E))
        self.size_var = tk.StringVar(value="2")
        ttk.Radiobutton(size_frame, text="小图 (4英寸宽)", variable=self.size_var, value="1").grid(row=0, column=0, padx=20)
        ttk.Radiobutton(size_frame, text="中等 (6英寸宽)", variable=self.size_var, value="2").grid(row=0, column=1, padx=20)
        ttk.Radiobutton(size_frame, text="大图 (8英寸宽)", variable=self.size_var, value="3").grid(row=0, column=2, padx=20)
        
        # 增强图片清晰度选项
        self.enhance_var = tk.BooleanVar(value=False)
        enhance_check = ttk.Checkbutton(main_frame, text="增强图片清晰度", variable=self.enhance_var)
        enhance_check.grid(row=6, column=0, pady=10)
        
        # 保留增强后的图片选项
        self.keep_var = tk.BooleanVar(value=False)
        keep_check = ttk.Checkbutton(main_frame, text="保留增强后的图片", variable=self.keep_var)
        keep_check.grid(row=7, column=0, pady=10)
        
        # 处理按钮：改为调用 start_processing，点击后按钮隐藏，进度条显示更新
        self.process_button = ttk.Button(main_frame, text="开始处理", command=self.start_processing)
        self.process_button.grid(row=8, column=0, pady=20)
        
        # 进度条（设置 maximum 为 100）
        self.progress = ttk.Progressbar(main_frame, length=400, mode='determinate', maximum=100)
        self.progress.grid(row=9, column=0, pady=10)
        
        # 状态标签（用于显示详细进度信息）
        self.status_label = ttk.Label(main_frame, text="", font=('Helvetica', 12))
        self.status_label.grid(row=10, column=0, pady=10)
        
        # 配置列权重
        main_frame.columnconfigure(0, weight=1)
        
        # macOS特定的菜单栏
        menubar = tk.Menu(root)
        root.config(menu=menubar)
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="文件", menu=file_menu)
        file_menu.add_command(label="选择PDF文件", command=self.select_pdf_file)
        file_menu.add_separator()
        file_menu.add_command(label="退出", command=root.quit)
        
        # 创建一个线程安全的队列用于传递UI更新信息
        self.ui_queue = queue.Queue()
        self.poll_ui_queue()

    def poll_ui_queue(self):
        """定时检查队列，更新UI"""
        try:
            while True:
                message, progress_value = self.ui_queue.get_nowait()
                self.status_label.config(text=message)
                if progress_value is not None:
                    self.progress['value'] = int(progress_value)
        except queue.Empty:
            pass
        self.root.after(100, self.poll_ui_queue)

    def update_status(self, message, progress_value=None):
        """将更新任务放入队列，由主线程处理"""
        self.ui_queue.put((message, progress_value))

    def select_pdf_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file_path:
            self.folder_label.config(text=file_path)
            self.pdf_file = file_path

    def start_processing(self):
        # 隐藏“开始处理”按钮
        self.process_button.grid_forget()
        threading.Thread(target=self.process_images, daemon=True).start()

    def process_images(self):
        file_path = self.folder_label.cget("text")
        if file_path == "未选择文件":
            messagebox.showerror("错误", "请先选择PDF文件！")
            self.root.after(0, lambda: self.process_button.grid(row=8, column=0, pady=20))
            return
        try:
            self.update_status("正在转换PDF为图片...", 0)
            temp_folder = os.path.join(os.path.dirname(file_path), "temp_images")
            if not os.path.exists(temp_folder):
                os.makedirs(temp_folder)
            image_files = pdf_to_images(file_path, temp_folder)
            if not image_files:
                return
            # 计算总步骤数：PDF转换、图片增强、图片重命名、文档创建（每张图片多步骤）、完成状态
            total_steps = 1
            if self.enhance_var.get():
                total_steps += len(image_files)
            total_steps += len(image_files)
            total_steps += len(image_files) * 3
            total_steps += 1
            step = 1
            total_images = len(image_files)
            self.update_status(f"共转换出 {total_images} 张图片", step / total_steps * 100)
            
            # 如果需要增强图片，逐张处理并更新进度
            enhanced_files = []
            if self.enhance_var.get():
                enhance_level = 1.5
                for i, img_path in enumerate(image_files, start=1):
                    enhanced_path = os.path.join(
                        temp_folder,
                        os.path.basename(img_path).replace("temp_page_", "enhanced_page_")
                    )
                    enhance_image(img_path, enhanced_path, enhance_level)
                    enhanced_files.append(enhanced_path)
                    if not self.keep_var.get():
                        os.remove(img_path)
                    step += 1
                    # 此处更新的进度值按比例计算，可根据实际情况调整
                    self.update_status(f"增强图片 {i}/{total_images}", 10 + i/total_images*10)
                image_files = enhanced_files

            renamed_files = self.rename_images(temp_folder, image_files)
            if not renamed_files:
                return
            step += 1
            self.update_status("开始创建文档...", step / total_steps * 100)
            self.create_image_document(temp_folder, step, total_steps)
            step += 1
            if not self.keep_var.get():
                shutil.rmtree(temp_folder, ignore_errors=True)
            self.update_status("处理完成！", step / total_steps * 100)
            # 恢复“开始处理”按钮
            self.root.after(0, lambda: self.process_button.grid(row=8, column=0, pady=20))
            messagebox.showinfo("完成", "文档已生成完成！")
        except Exception as e:
            messagebox.showerror("错误", f"处理过程中出错: {str(e)}")
            logging.error(f"处理过程中出错: {str(e)}")
            self.root.after(0, lambda: self.process_button.grid(row=8, column=0, pady=20))

    def extract_number_from_filename(self, filename):
        match = re.search(r'_(\d{2})(?:\.|$)', filename)
        if match:
            return int(match.group(1))
        return 9999

    def rename_images(self, directory, image_files):
        self.update_status("正在重命名图片...", 20)
        try:
            renamed_files = []
            for index, img_path in enumerate(image_files, 1):
                new_name = f"幻灯片 {index:02d}{os.path.splitext(img_path)[1]}"
                new_path = os.path.join(directory, new_name)
                os.rename(img_path, new_path)
                renamed_files.append(new_name)
                self.update_status(f"重命名图片 {index}/{len(image_files)}", 20 + index/len(image_files)*10)
            return renamed_files
        except Exception as e:
            messagebox.showerror("错误", f"重命名过程中出错: {str(e)}")
            logging.error(f"重命名过程中出错: {str(e)}")
            return None

    def create_image_document(self, input_directory, start_step, total_steps):
        doc = Document()
        image_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
        img_width = {'1': 4, '2': 6, '3': 8}[self.size_var.get()]
        def get_slide_number(filename):
            match = re.search(r'幻灯片 (\d+)', filename)
            if match:
                return int(match.group(1))
            return 9999
        image_files = [f for f in os.listdir(input_directory) if f.lower().endswith(image_extensions)]
        image_files.sort(key=get_slide_number)
        step = start_step
        total_images = len(image_files)
        self.update_status(f"开始处理 {total_images} 张图片...", 40)
        logging.info(f"共需要处理 {total_images} 张图片")
        
        for index, filename in enumerate(image_files, 1):
            image_path = os.path.join(input_directory, filename)
            try:
                file_name_without_ext = os.path.splitext(filename)[0]
                paragraph = doc.add_paragraph(file_name_without_ext)
                paragraph._p.get_or_add_pPr().append(
                    parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFC000"/>')
                )
                step += 1
                self.update_status(f"插入文件名 {index}/{total_images}", step / total_steps * 100)
                doc.add_picture(image_path, width=Inches(img_width))
                step += 1
                self.update_status(f"插入图片 {index}/{total_images}", step / total_steps * 100)
                self.update_status(f"图片 {index}/{total_images}: 正在进行OCR识别...", step / total_steps * 100)
                logging.info(f"图片 {index}/{total_images}: 开始OCR识别 {filename}")
                original_text = ocr_text(image_path)
                step += 1
                self.update_status(f"OCR识别完成 {index}/{total_images}", step / total_steps * 100)
                if original_text:
                    self.update_status(f"图片 {index}/{total_images}: 正在翻译...", step / total_steps * 100)
                    logging.info(f"图片 {index}/{total_images}: 开始翻译 {filename}")
                    translated_text = translate_text_in_chunks(original_text, self.language_var.get())
                    step += 1
                    self.update_status(f"翻译完成 {index}/{total_images}", step / total_steps * 100)
                    if translated_text:
                        doc.add_paragraph(f"原文: {original_text}")
                        doc.add_paragraph(f"翻译: {translated_text}")
                    else:
                        doc.add_paragraph(f"原文: {original_text}")
                        doc.add_paragraph("翻译: （翻译失败）")
                else:
                    doc.add_paragraph("原文: （无识别内容）")
                    doc.add_paragraph("翻译: （无识别内容）")
                    step += 1
                    self.update_status(f"处理无识别内容 {index}/{total_images}", step / total_steps * 100)
                doc.add_paragraph()
                step += 1
                self.update_status(f"处理完成 {index}/{total_images}", step / total_steps * 100)
                logging.info(f"图片 {index}/{total_images}: 处理完成 {filename}")
            except Exception as e:
                messagebox.showerror("错误", f"处理图片 {filename} 时出错: {str(e)}")
                logging.error(f"处理图片 {filename} 时出错: {str(e)}")
                return
        output_file = os.path.join(
            os.path.dirname(input_directory),
            f"{os.path.basename(self.pdf_file)}_转换结果.docx"
        )
        doc.save(output_file)
        logging.info(f"文档已保存到: {output_file}")

def main():
    root = tk.Tk()
    app = ImageProcessorApp(root)
    if sys.platform == 'darwin':
        root.createcommand('tk::mac::ReopenApplication', root.lift)
    try:
        root.mainloop()
    finally:
        if hasattr(ocr_text, "ocr_engine"):
            del ocr_text.ocr_engine
        if sys.platform == 'darwin':
            subprocess.run(['purge'], check=True)

if __name__ == "__main__":
    main()
